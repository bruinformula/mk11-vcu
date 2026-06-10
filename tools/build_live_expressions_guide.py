from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "STM32CubeIDE_Live_Expressions_Operation_Guide.docx"

BLACK = "000000"
GRAY = "555555"
BORDER = "DADCE0"
LIGHT = "F8F9FA"
WARN = "FFF4E5"
CURRENT_NUM_ID = None


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER, size="4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    cant_split.set(qn("w:val"), "true")
    tr_pr.append(cant_split)


def set_table_fixed_layout(table):
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")


def set_col_width(cell, width_inches):
    cell.width = Inches(width_inches)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    tc_w.set(qn("w:w"), str(int(width_inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def add_table(doc, headers, rows, widths, font_size=8.5):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_fixed_layout(table)
    set_table_borders(table)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    prevent_row_split(hdr)
    for i, (text, width) in enumerate(zip(headers, widths)):
        set_col_width(hdr.cells[i], width)
        set_cell_shading(hdr.cells[i], LIGHT)
        set_cell_margins(hdr.cells[i])
        hdr.cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = hdr.cells[i].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        r.bold = True
        r.font.name = "Arial"
        r.font.size = Pt(font_size)
    for row_data in rows:
        row = table.add_row()
        prevent_row_split(row)
        for i, (text, width) in enumerate(zip(row_data, widths)):
            set_col_width(row.cells[i], width)
            set_cell_margins(row.cells[i])
            row.cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = row.cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(text))
            r.font.name = "Arial"
            r.font.size = Pt(font_size)
            if i == 0:
                r.font.name = "Arial"
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.add_run(text)
    return p


def begin_numbered_list(doc):
    global CURRENT_NUM_ID
    numbering = doc.part.numbering_part.element
    nums = numbering.findall(qn("w:num"))
    next_num_id = max(int(n.get(qn("w:numId"))) for n in nums) + 1

    style_num_id = doc.styles["List Number"].element.pPr.numPr.numId.val
    style_num = numbering.find(f".//{qn('w:num')}[@{qn('w:numId')}='{style_num_id}']")
    abstract_num_id = style_num.find(qn("w:abstractNumId")).get(qn("w:val"))

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(next_num_id))
    abstract = OxmlElement("w:abstractNumId")
    abstract.set(qn("w:val"), abstract_num_id)
    num.append(abstract)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:startOverride")
    start.set(qn("w:val"), "1")
    override.append(start)
    num.append(override)
    numbering.append(num)
    CURRENT_NUM_ID = next_num_id


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p_pr = p._p.get_or_add_pPr()
    num_pr = p_pr.get_or_add_numPr()
    ilvl = num_pr.get_or_add_ilvl()
    ilvl.set(qn("w:val"), "0")
    num_id = num_pr.get_or_add_numId()
    num_id.set(qn("w:val"), str(CURRENT_NUM_ID))
    p.add_run(text)
    return p


def add_note(doc, label, text, fill=LIGHT):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_fixed_layout(table)
    set_table_borders(table, color=BORDER)
    cell = table.cell(0, 0)
    set_col_width(cell, 6.5)
    set_cell_margins(cell, top=120, start=160, bottom=120, end=160)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(f"{label}: ")
    r.bold = True
    p.add_run(text)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def configure_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    for style_name, size, color, before, after in (
        ("Heading 1", 20, BLACK, 20, 6),
        ("Heading 2", 16, BLACK, 18, 6),
        ("Heading 3", 14, "434343", 16, 4),
    ):
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = False
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Bullet 2", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.15


doc = Document()
configure_styles(doc)

title = doc.add_paragraph()
title.paragraph_format.space_before = Pt(0)
title.paragraph_format.space_after = Pt(3)
r = title.add_run("STM32CubeIDE Live Expressions Operation Guide")
r.font.name = "Arial"
r.font.size = Pt(26)
r.font.bold = False
r.font.color.rgb = RGBColor(0, 0, 0)

subtitle = doc.add_paragraph()
subtitle.paragraph_format.space_after = Pt(14)
r = subtitle.add_run("MK11 VCU firmware | Operator-focused monitoring and function reference")
r.font.name = "Arial"
r.font.size = Pt(12)
r.font.color.rgb = RGBColor.from_string(GRAY)

doc.add_paragraph(
    "This guide is for a test operator, controls engineer, or team member who needs to run the "
    "vehicle control unit (VCU) firmware and observe its important internal values without first "
    "learning the entire codebase. It covers application globals and file-local static variables "
    "that are useful in STM32CubeIDE Live Expressions, plus the functions that update them."
)

add_note(
    doc,
    "Safety",
    "Treat Live Expressions as read-only. Do not change state, torque, relay, plausibility, CAN, "
    "timer, or peripheral values while the tractive system is energized. Editing a value can bypass "
    "normal safety logic or produce an unsafe command.",
    WARN,
)

doc.add_heading("1. Quick Start", level=1)
begin_numbered_list(doc)
add_number(doc, "Build the Debug configuration and start a normal debug session on the VCU.")
add_number(doc, "Open Window > Show View > Other > Debug > Live Expressions.")
add_number(doc, "Use Add new expression and enter the exact expressions from the Recommended Watch Set below.")
add_number(doc, "Resume the target. Live Expressions refreshes while the MCU runs; ordinary Expressions may require a halt.")
add_number(doc, "Verify the state sequence first, then pedals, inverter feedback, cooling, CAN counters, and audio as needed.")

add_note(
    doc,
    "Static variables",
    "File-local statics are present in the Debug build and are included here. If CubeIDE cannot "
    "resolve a static by its plain name, halt inside its source file/function and add it from the "
    "Variables or Expressions view. Link-time garbage collection can remove an unused static or "
    "unused function entirely.",
)

doc.add_heading("2. Recommended Watch Set", level=1)
doc.add_paragraph(
    "Start with this compact dashboard. Array elements and structure fields are entered exactly as shown."
)
watch_rows = [
    ("vcu_state", "Overall operating state", "0 Idle; 1 Precharged; 2 Drive; 3 CAN fault"),
    ("precharge_state", "Precharge progress", "0 Idle; 1 Waiting; 2 Success; 3 Failure; 4 Timeout"),
    ("precharge_response_received", "BMS reply received", "0 no; 1 yes"),
    ("ADC_VAL[0]", "APPS1 raw ADC", "Nominal released ~1280; pressed ~295"),
    ("ADC_VAL[1]", "APPS2 raw ADC", "Nominal released ~1820; pressed ~435"),
    ("ADC_VAL[2]", "Brake sensor raw ADC", "Brake active threshold >900"),
    ("pedal_percents[0]", "APPS1 normalized", "Approximately 0.0 released to 1.0 pressed"),
    ("pedal_percents[1]", "APPS2 normalized", "Approximately 0.0 released to 1.0 pressed"),
    ("pedal_percents[2]", "Brake normalized", "Approximately 0.0 to 1.0"),
    ("plausibility_checks.apps1_invalid", "APPS1 range fault", "0 expected; 1 torque is cut"),
    ("plausibility_checks.apps2_invalid", "APPS2 range fault", "0 expected; 1 torque is cut"),
    ("plausibility_checks.apps_plausible", "APPS agreement", "1 expected"),
    ("plausibility_checks.crosscheck_plausible", "Accelerator/brake cross-check", "1 expected"),
    ("requestedTorque", "Final filtered torque request", "Nm; positive drive, negative regen"),
    ("inverter_diagnostics.inverter_voltage", "DC bus voltage", "V"),
    ("inverter_diagnostics.inverter_rpm", "Motor speed", "RPM, absolute value"),
    ("inverter_diagnostics.inverter_carspeed", "Estimated vehicle speed", "mph"),
    ("bms_avg_temp", "BMS average temperature", "deg C"),
    ("inv_avg_temp", "Average inverter phase temperature", "deg C"),
    ("last_tractive_fan_pwm", "Tractive fan command", "0-100 percent; idle command 10"),
    ("last_tractive_pump_pwm", "Tractive pump command", "0-100 percent; idle command 50"),
    ("last_accy_fan_pwm", "Accumulator fan command", "0-75 percent in Drive; idle command 10"),
    ("RxHeader1.Identifier", "Most recent drive CAN ID", "Hex CAN identifier"),
    ("fdcan_rx_count", "Total accepted CAN messages", "Should increase with bus traffic"),
    ("fdcan_rx_error_count", "CAN receive failures", "Should remain 0"),
]
add_table(doc, ["Expression", "Meaning", "Normal interpretation"], watch_rows, [2.40, 1.85, 2.25], 8.25)

doc.add_heading("3. State and Startup Sequence", level=1)
doc.add_paragraph(
    "The VCU state machine is the first place to look. A normal run progresses from Idle to "
    "Precharged to Drive. A shutdown-power-lost CAN message resets the system to Idle."
)
add_table(
    doc,
    ["Value", "State", "What the operator should see"],
    [
        ("0", "VCU_IDLE", "Startup/reset state. Torque disabled. Precharge button is accepted."),
        ("1", "VCU_PRECHARGED", "BMS precharge succeeded. Ready-to-drive button is accepted."),
        ("2", "VCU_DRIVE", "Pedal processing and inverter torque commands are active."),
        ("3", "VCU_CAN_FAULT", "Precharge timed out. Firmware enters Error_Handler and stops normal operation."),
    ],
    [0.65, 1.55, 4.30],
)
add_table(
    doc,
    ["Value", "Precharge state", "Interpretation"],
    [
        ("0", "PRECHARGE_IDLE", "No request active."),
        ("1", "PRECHARGE_WAITING", "Request sent; six-second timeout timer is running."),
        ("2", "PRECHARGE_SUCCESS", "BMS returned success."),
        ("3", "PRECHARGE_FAILURE", "BMS returned a failure response."),
        ("4", "PRECHARGE_TIMEOUT", "No response before timeout; VCU enters CAN fault."),
    ],
    [0.65, 1.80, 4.05],
)

doc.add_heading("4. Pedals, Plausibility, and Torque", level=1)
doc.add_paragraph(
    "ADC3 continuously samples APPS1, APPS2, and the brake sensor through DMA. Pedal percentages "
    "and torque are updated only in VCU_DRIVE and after the 150 ms inverter lockout period."
)
pedal_rows = [
    ("ADC_VAL", "uint16_t[3], global", "[0] APPS1; [1] APPS2; [2] brake. Raw 12-bit ADC samples."),
    ("voltage_values", "float[3], global", "Converted sensor voltage using raw/4095 x 3.3 V."),
    ("pedal_percents", "float[3], global", "Normalized sensor travel. Values can go outside 0-1 if calibration/range is exceeded."),
    ("plausibility_checks", "struct, global", "Range validity, APPS agreement, and accelerator/brake cross-check flags."),
    ("millis_since_apps_implausible", "uint32_t, static", "Tick when APPS disagreement first became active."),
    ("filtered_target_torque", "float, static", "Low-pass-filtered target before the final slew-rate limiter."),
    ("last_torque_calc_time", "uint32_t, static", "Tick used to calculate torque filter time step."),
    ("requestedTorque", "float, global", "Final command in Nm before multiplication by 10 for the inverter CAN payload."),
    ("inverter_lockout_start", "uint32_t, global", "Tick when Drive begins; torque processing waits 150 ms."),
    ("inverter_tx_rate_limiter", "uint32_t, global", "Tick of last inverter torque frame; nominal minimum interval 10 ms."),
]
add_table(doc, ["Variable", "Type/scope", "Operational meaning"], pedal_rows, [1.85, 1.55, 3.10], 8.5)

doc.add_heading("Torque behavior", level=2)
add_bullet(doc, "Any APPS range fault, failed APPS/brake cross-check, or sustained APPS disagreement commands zero torque.")
add_bullet(doc, "At or above 5% accelerator, torque scales toward 115 Nm.")
add_bullet(doc, "Below 5% accelerator, the code may request as much as -30 Nm regen.")
add_bullet(doc, "Regen is cancelled below 5 mph or when normalized brake travel exceeds 5%.")
add_bullet(doc, "The request passes through a 0.05 s low-pass filter and a 500 Nm/s slew-rate limit.")

doc.add_heading("5. Inverter and CAN Monitoring", level=1)
can_rows = [
    ("inverter_diagnostics.inverter_voltage", "Decoded from CAN 0x0A7 bytes 0-1", "V; raw signed value x 0.1"),
    ("inverter_diagnostics.inverter_rpm", "Decoded from CAN 0x0A5 bytes 2-3", "Absolute RPM"),
    ("inverter_diagnostics.inverter_carspeed", "Calculated from RPM", "Estimated mph"),
    ("RxHeader1.Identifier", "Latest FDCAN1 header", "Use hexadecimal display to identify the frame"),
    ("RxData1[0] ... RxData1[7]", "Latest FDCAN1 payload", "Raw bytes used by the dispatch handler"),
    ("RxHeader2 / RxData2", "Latest FDCAN2 frame", "DAQ bus receive storage; handler is currently empty"),
    ("fdcan1_irq_count", "FDCAN1 ISR entry count", "Should increase when FDCAN1 interrupts occur"),
    ("fdcan1_debug_cb", "FDCAN1 receive callback count", "Should track valid FIFO0 callbacks"),
    ("fdcan2_debug_cb", "FDCAN2 receive callback count", "Should track FIFO1 callbacks"),
    ("fdcan_rx_count", "Successful receive count", "Increments after HAL_FDCAN_GetRxMessage succeeds"),
    ("fdcan_rx_error_count", "Receive error count", "Expected 0"),
    ("hfdcan1 / hfdcan2", "HAL peripheral handles", "Advanced diagnosis only; inspect ErrorCode, State, and Instance"),
]
add_table(doc, ["Expression", "Source", "Use"], can_rows, [2.20, 2.25, 2.05], 8.4)

doc.add_heading("Relevant CAN IDs", level=2)
add_table(
    doc,
    ["ID", "Direction", "Purpose"],
    [
        ("0x0C0", "VCU -> inverter", "Torque request, direction, inverter enable"),
        ("0x500", "VCU -> DAQ", "VCU diagnostics"),
        ("0x501", "VCU -> cooling controller", "Cooling enable and PWM commands"),
        ("0x6B3", "VCU -> BMS", "Precharge request"),
        ("0x0A0", "Inverter -> VCU", "Phase temperatures"),
        ("0x0A5", "Inverter -> VCU", "RPM"),
        ("0x0A7", "Inverter -> VCU", "DC bus voltage"),
        ("0x6B1", "BMS -> VCU", "Average temperature"),
        ("0x6B4", "BMS -> VCU", "Precharge response"),
        ("0x6B5", "System -> VCU", "Shutdown power lost; triggers resetVCU"),
    ],
    [0.9, 1.75, 3.85],
)

doc.add_heading("6. Cooling Monitoring", level=1)
cooling_rows = [
    ("bms_avg_temp", "deg C", "BMS raw temperature / 100"),
    ("phase_a_temp", "deg C", "Inverter phase A raw / 10"),
    ("phase_b_temp", "deg C", "Inverter phase B raw / 10"),
    ("phase_c_temp", "deg C", "Inverter phase C raw / 10"),
    ("inv_avg_temp", "deg C", "Average of three inverter phase temperatures"),
    ("last_tractive_fan_pwm", "%", "10 outside Drive; 0-100 across 20-50 C in Drive"),
    ("last_tractive_pump_pwm", "%", "50 outside Drive; 0-100 across 20-50 C in Drive"),
    ("last_accy_fan_pwm", "%", "10 outside Drive; 0-75 across 20-55 C in Drive"),
    ("cooling_cmd_df.data.cooling_en", "0/1", "Set to 1 when the 2 Hz command is assembled"),
    ("last_cooling_cmd_send", "HAL tick ms", "Last successful cooling command transmission"),
]
add_table(doc, ["Variable", "Unit", "Interpretation"], cooling_rows, [2.25, 1.00, 3.25], 8.5)

doc.add_heading("7. Buttons, Timers, and Ready-to-Drive Audio", level=1)
button_rows = [
    ("prchg_button_pressed", "Precharge button level captured by EXTI", "Must remain true until TIM1 debounce/hold callback."),
    ("rtd_button_pressed", "Ready-to-drive button level captured by EXTI", "Must remain true until TIM2 callback."),
    ("prchg_debug", "Precharge edge counter", "Increments on each EXTI callback for the precharge pin."),
    ("rtd_debug", "RTD edge counter", "Increments on each EXTI callback for the RTD pin."),
    ("sound_success", "Result of ready-to-drive playback", "Must be true before Drive entry continues."),
    ("debug.playback_kind", "Audio mode", "1 test tone; 2 ready-to-drive sound."),
    ("debug.playback_started", "Playback attempts", "Increments when playback begins."),
    ("debug.playback_finished", "Completed playback count", "Increments on successful completion."),
    ("debug.i2s_state", "HAL I2S state snapshot", "Use with debug.i2s_error when audio fails."),
    ("debug.last_i2s_status", "Last HAL transmit result", "HAL_OK is 0."),
    ("debug.i2s_error_cb_hit", "I2S error callback count", "Expected 0."),
    ("wavPos / halfwordCount", "Playback progress", "wavPos advances toward halfwordCount."),
    ("waveFinished", "Playback completion flag", "1 after completion or a blocking transmit failure."),
]
add_table(doc, ["Expression", "Meaning", "Expected use"], button_rows, [2.15, 2.20, 2.15], 8.25)

doc.add_heading("8. Function Reference", level=1)
doc.add_paragraph(
    "These functions are listed as execution landmarks: they explain where values change and are "
    "good breakpoint locations. Do not call safety-critical functions manually from the debugger."
)
function_rows = [
    ("main", "Startup", "Initializes peripherals, CAN filters, message headers, ADC DMA; loops sending cooling and diagnostics."),
    ("HAL_ADC_ConvCpltCallback", "Pedals/torque", "Converts ADC values, runs plausibility checks, calculates and sends torque in Drive."),
    ("validateAPPS", "Pedals/torque", "Checks each APPS raw channel against calibrated range plus overshoot buffer."),
    ("checkAPPS_Plausibility", "Pedals/torque", "Checks APPS1 versus APPS2 agreement and records disagreement time."),
    ("checkAPPS_BSE_Crosscheck", "Pedals/torque", "Latches accelerator/brake conflict until accelerator returns below 5%."),
    ("calculateTorqueRequest", "Pedals/torque", "Selects drive/regen/zero torque, then filters and slew-limits it."),
    ("sendTorqueRequest", "Inverter TX", "Builds and queues CAN 0x0C0. Has direct actuator consequences."),
    ("processInverter_Voltage", "Inverter RX", "Decodes DC bus voltage from the latest RxData1 payload."),
    ("processInverter_RPM", "Inverter RX", "Decodes absolute RPM and computes estimated vehicle speed."),
    ("HAL_GPIO_EXTI_Callback", "Buttons", "Tracks precharge and RTD button edges and starts/stops hold timers."),
    ("HAL_TIM_PeriodElapsedCallback", "State transitions", "Handles button hold completion, precharge timeout, audio, and Drive entry."),
    ("sendPrechargeRequest", "Precharge", "Sends BMS request, sets Waiting state, and starts timeout timer."),
    ("processPrechargeResponse", "Precharge", "Converts BMS response byte into Success or Failure."),
    ("checkPrechargeStatus", "Precharge", "On timeout, advances to Precharged or enters CAN fault/Error_Handler."),
    ("resetVCU", "Reset/safety", "Stops activity, clears state, sends inverter-disable frames, and restarts pedal ADC."),
    ("FDCAN1_Rx_Handler", "CAN dispatch", "Routes known drive-bus IDs to precharge, inverter, cooling, or reset processing."),
    ("HAL_FDCAN_RxFifo0Callback", "CAN receive", "Reads FIFO0, updates counters, then calls FDCAN1_Rx_Handler."),
    ("processBMS_Temp", "Cooling", "Decodes BMS temperature and updates accumulator fan command."),
    ("processInverter_Temp", "Cooling", "Decodes phase temperatures and updates tractive fan/pump commands."),
    ("sendCoolingCmd", "Cooling", "Queues CAN 0x501 at up to 2 Hz."),
    ("sendVCUDiagnostics", "Diagnostics", "Queues CAN 0x500 at up to 20 Hz."),
    ("playReadyToDriveSound", "Audio", "Blocking RTD sound playback; returns success/failure."),
    ("captureI2SDebugState", "Audio", "Snapshots I2S state, error, clock, and registers into debug."),
    ("Error_Handler", "Fatal fault", "Sends inverter-disable frames, stops I2S/ADC, disables interrupts, and loops forever."),
]
add_table(doc, ["Function", "Area", "What it does"], function_rows, [2.25, 1.25, 3.00], 8.15)

doc.add_heading("9. Operator Test Flows", level=1)
doc.add_heading("Normal startup to Drive", level=2)
begin_numbered_list(doc)
add_number(doc, "Confirm vcu_state = 0 and precharge_state = 0.")
add_number(doc, "Press and hold precharge. Confirm prchg_button_pressed becomes 1, then precharge_state becomes 1.")
add_number(doc, "Confirm precharge_response_received becomes 1 and precharge_state becomes 2; after timeout handling, vcu_state becomes 1.")
add_number(doc, "Hold the brake so ADC_VAL[2] > 900, then press and hold RTD.")
add_number(doc, "Confirm sound_success = 1, inverter_lockout_start updates, and vcu_state becomes 2.")
add_number(doc, "With pedals released, verify plausibility flags are healthy and requestedTorque is near 0.")

doc.add_heading("Pedal plausibility check", level=2)
begin_numbered_list(doc)
add_number(doc, "Observe both raw APPS channels and normalized percentages while pressing the pedal slowly.")
add_number(doc, "Confirm apps1_invalid = 0, apps2_invalid = 0, and apps_plausible = 1.")
add_number(doc, "Confirm the two normalized APPS values remain within 0.10 of each other.")
add_number(doc, "Apply brake and accelerator together above the thresholds; crosscheck_plausible should become 0 and requestedTorque should become 0.")
add_number(doc, "Release accelerator below 5%; crosscheck_plausible should restore to 1.")

doc.add_heading("CAN diagnosis", level=2)
begin_numbered_list(doc)
add_number(doc, "Check fdcan1_irq_count and fdcan1_debug_cb. If neither changes, inspect wiring, transceiver power, bus termination, and FDCAN configuration.")
add_number(doc, "If the IRQ count changes but callback count does not, inspect notification/filter/FIFO configuration.")
add_number(doc, "If callback count changes but fdcan_rx_count does not, inspect fdcan_rx_error_count and hfdcan1.ErrorCode.")
add_number(doc, "Use RxHeader1.Identifier and RxData1 bytes to confirm expected source messages.")

doc.add_heading("10. Low-Level and Usually Optional Symbols", level=1)
doc.add_paragraph(
    "These are valid globals/statics in the Debug image but are normally useful only when a subsystem "
    "has already been isolated. Large buffers can make Live Expressions noisy and should not be expanded continuously."
)
optional_rows = [
    ("hadc3, hdma_adc3", "ADC/DMA HAL handles", "Inspect State/ErrorCode for ADC acquisition faults."),
    ("hi2s2, hdma_spi2_tx", "I2S/DMA HAL handles", "Inspect State/ErrorCode for sound faults."),
    ("htim1, htim2, htim3", "Timer HAL handles", "Precharge button, RTD button, and precharge timeout timers."),
    ("Inverter_TxHeader / Inverter_TxData", "Static CAN TX storage", "Raw inverter command construction."),
    ("PRCHG_TxHeader / PRCHG_TxData", "Static CAN TX storage", "Raw BMS precharge request."),
    ("CoolingCmd_TxHeader / cooling_cmd_df", "Static CAN TX storage", "Raw cooling command."),
    ("VCU_Diagnostics_TxHeader / VCU_Diagnostics_TxData", "Static CAN TX storage", "Raw DAQ diagnostics frame."),
    ("wavePCM", "Static audio pointer", "Current PCM source."),
    ("scaledPlaybackBuffer", "Static 4096-halfword buffer", "Scaled RTD audio chunk; do not expand live unless needed."),
    ("SystemCoreClock / SystemD2Clock", "CMSIS clock globals", "Clock-tree diagnosis only."),
    ("uwTick", "HAL millisecond tick", "Confirms SysTick is advancing."),
]
add_table(doc, ["Symbols", "Category", "When useful"], optional_rows, [2.35, 1.70, 2.45], 8.35)

doc.add_heading("11. Troubleshooting Live Expressions", level=1)
add_bullet(doc, "Expression says “not available”: confirm a Debug build with -g3 and optimization -O0, then rebuild and restart the debug session.")
add_bullet(doc, "Static cannot be found: halt in the owning source file, add the variable from the normal Expressions view, or set a breakpoint in a function that uses it.")
add_bullet(doc, "Value never changes: confirm the function that updates it is executing. Use the function reference and counters to choose a breakpoint.")
add_bullet(doc, "Value looks stale: Live Expressions refresh can be slower than a fast ISR. Watch the processed value or counter rather than every raw transition.")
add_bullet(doc, "Array/struct is hard to read: add individual fields such as pedal_percents[0] or plausibility_checks.apps_plausible.")
add_bullet(doc, "Release build behaves differently: optimized variables may be folded, relocated, or removed. Use the Debug configuration for this guide.")

doc.add_heading("12. Source and Scope", level=1)
doc.add_paragraph(
    "Prepared from the MK11 VCU workspace on June 9, 2026. Application behavior was traced through "
    "Core/Src/main.c, motor_control.c, vcu_state.c, prchg.c, fdcan.c, cooling.c, audio.c, adc.c, "
    "tim.c, i2s.c, and stm32h7xx_it.c, with symbols cross-checked against Debug/mk11-vcu.elf. "
    "Generated STM32 HAL/CMSIS internals and immutable audio sample data are intentionally excluded "
    "except where they provide useful operational diagnostics."
)
doc.add_paragraph(
    "The source tree contained an uncommitted whitespace-only change in main.c at preparation time. "
    "This guide describes the current source behavior and the current Debug symbol set."
)

doc.save(OUT)
print(OUT)
